from docx import Document
 
 
def markdown_to_docx(text):
    document = Document()
    lines = text.split('\n')
    name = lines[0]
    document.add_heading(name, 0)
    t = False
    for line in lines[1:]:
        if line:
            t = False
            if line[:7].count('#'):
                n = line[:7].count('#')
                p = document.add_heading(level=n)
                line = line[n + 1:]
            elif line[:2] in ('- ', '* ', '+ '):
                p = document.add_paragraph(style='List Bullet')
                line = line[2:]
            elif line[0].isdigit() and line[1] == '.':
                p = document.add_paragraph(style='List Number')
                line = line[3:]
            else:
                p = document.add_paragraph()
            c = ''
            if line.count('*') % 2 == 0 and line.count('_') % 2 > 0:
                c = '*'
            elif line.count('*') % 2 > 0 and line.count('_') % 2 == 0:
                c = '_'
            elif line.count('*') % 2 == 0 == line.count('_') % 2:
                c = '*_'
            if c:
                t = 0
                italic, bold = False, False
                for i in line:
                    if i in c:
                        t += 1
                    else:
                        if t == 1:
                            italic = not italic
                        elif t == 2:
                            bold = not bold
                        elif t == 3:
                            italic, bold = not italic, not bold
                        t = 0
                        run = p.add_run(i)
                        run.italic, run.bold = italic, bold
            else:
                p.add_run(line)
        else:
            if t:
                document.add_paragraph()
            t = True
    document.save('res2.docx')

markdown_text = """
Документ по Python

# Введение
Python — это язык программирования, который поддерживает несколько парадигм.

## Возможности
- Простота 
- Читаемость
- Широкая библиотека

1. Язык Python работает практически на всех известных платформах — от карманных компьютеров и смартфонов до серверов сети
2. Python — один из самых популярных в мире языков программирования.
3. Название Python произошло не от рептилии.

**Надеюсь, вам это понравится!**
_Удачи!_
***Всегда учитесь!***
"""

markdown_to_docx(markdown_text)